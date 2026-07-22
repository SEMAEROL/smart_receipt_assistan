"""
Word Rapor Üretici Modülü

Bu modül python-docx kullanarak Receipt nesnelerini
profesyonel bir Word raporuna (.docx) dönüştürür.
"""

import logging
from typing import List, Dict, Optional
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.models.receipt import Receipt, BatchReportSummary

# Logging konfigürasyonu
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocxReportGenerator:
    """
    Word (.docx) formatında profesyonel raporlar üreten sınıf.
    
    python-docx kütüphanesi kullanarak:
    - Başlık ve genel bilgiler
    - Kategori bazlı harcama dağılımı tablosu
    - Fiş bazlı detay ve kalem tabloları
    oluşturur.
    
    Attributes:
        title (str): Rapor başlığı
        author (str): Rapor yazarı
        doc (Document): python-docx Document nesnesi
    """

    def __init__(self, title: str = "Harcama Raporu", author: str = "Smart Receipt Assistant"):
        """
        DocxReportGenerator başlatıcısı.
        
        Args:
            title (str): Rapor başlığı
            author (str): Rapor yazarı
        """
        self.title = title
        self.author = author
        self.doc = Document()
        
        # Varsayılan yazı tipi
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        logger.info(f"DocxReportGenerator başlatıldı - Başlık: {title}")

    def _set_cell_background(self, cell, color: str):
        """
        Tablo hücresinin arka plan rengini ayarlar.
        
        Args:
            cell: python-docx hücre nesnesi
            color (str): RGB renk (örn: 'D3D3D3')
        """
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def _add_heading(self, text: str, level: int = 1) -> None:
        """
        Rapora başlık ekler.
        
        Args:
            text (str): Başlık metni
            level (int): Başlık seviyesi (1-3)
        """
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Renk ayarla
        if level == 1:
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 128)  # Koyu mavi
        
        logger.debug(f"Başlık eklendi (Level {level}): {text}")

    def _add_summary_section(self, summary: BatchReportSummary) -> None:
        """
        Genel özet bilgilerini rapora ekler.
        
        Args:
            summary (BatchReportSummary): Rapor özeti
        """
        logger.info("Özet bölümü oluşturuluyor...")

        self._add_heading("Genel Özet", level=2)

        # Özet bilgileri tablosu
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        # Başlık satırı
        header_cells = table.rows[0].cells
        header_cells[0].text = "Metrik"
        header_cells[1].text = "Değer"
        
        # Başlık formatlaması
        for cell in header_cells:
            self._set_cell_background(cell, 'D3D3D3')
            for run in cell.paragraphs[0].runs:
                run.font.bold = True

        # Veri satırları
        data = [
            ("Toplam Fiş Sayısı", str(summary.total_receipts)),
            ("Genel Toplam Tutar", f"₺ {summary.total_amount:,.2f}"),
            ("Toplam KDV", f"₺ {summary.total_tax:,.2f}"),
            ("Ortalama Fiş Tutarı", f"₺ {summary.get_average_receipt_amount():,.2f}"),
        ]

        for idx, (label, value) in enumerate(data, start=1):
            cells = table.rows[idx].cells
            cells[0].text = label
            cells[1].text = value
            
            # Başlık sütununu gri yap
            self._set_cell_background(cells[0], 'E8E8E8')

        # Ek bilgi
        self.doc.add_paragraph(
            f"Rapor Tarihi: {summary.generation_date.strftime('%d.%m.%Y %H:%M:%S')}"
        ).runs[0].italic = True

    def _add_category_distribution(self, summary: BatchReportSummary) -> None:
        """
        Kategori bazlı harcama dağılım tablosunu rapora ekler.
        
        Args:
            summary (BatchReportSummary): Rapor özeti
        """
        if not summary.category_totals:
            logger.warning("Kategori verisi yok, tablo atlanıyor")
            return

        logger.info("Kategori dağılım tablosu oluşturuluyor...")

        self._add_heading("Kategori Bazlı Harcama Dağılımı", level=2)

        # Tablo oluştur
        table = self.doc.add_table(rows=len(summary.category_totals) + 1, cols=3)
        table.style = 'Light Grid Accent 1'

        # Başlık satırı
        header_cells = table.rows[0].cells
        headers = ["Kategori", "Tutar (₺)", "Yüzde (%)"]
        for idx, header in enumerate(headers):
            header_cells[idx].text = header
            self._set_cell_background(header_cells[idx], 'D3D3D3')
            for run in header_cells[idx].paragraphs[0].runs:
                run.font.bold = True

        # Veri satırları
        for idx, (category, amount) in enumerate(sorted(summary.category_totals.items()), start=1):
            cells = table.rows[idx].cells
            percentage = summary.get_category_percentage(category)
            
            cells[0].text = category
            cells[1].text = f"₺ {amount:,.2f}"
            cells[2].text = f"%{percentage:.1f}"

        self.doc.add_paragraph()  # Boş satır

    def _add_receipt_details(self, receipts: List[Receipt]) -> None:
        """
        Fiş bazlı detay bilgilerini rapora ekler.
        Her fiş için ayrı bir tablo oluşturur.
        
        Args:
            receipts (List[Receipt]): Receipt nesnelerinin listesi
        """
        if not receipts:
            logger.warning("Fiş verisi yok, detaylar atlanıyor")
            return

        logger.info(f"Fiş detayları oluşturuluyor ({len(receipts)} fiş)...")

        self._add_heading("Fiş Detayları", level=2)

        for receipt in receipts:
            # Fiş başlığı
            fiş_heading = self.doc.add_heading(
                f"{receipt.merchant_name} - {receipt.date}",
                level=3
            )
            fiş_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Fiş meta bilgileri
            info_paragraph = self.doc.add_paragraph()
            info_text = (
                f"Dosya: {receipt.filename} | "
                f"Kategori: {receipt.category} | "
                f"Toplam: ₺ {receipt.total_amount:,.2f} | "
                f"KDV: ₺ {receipt.tax_amount:,.2f}"
            )
            info_paragraph.add_run(info_text).italic = True
            info_paragraph.space_before = Pt(6)
            info_paragraph.space_after = Pt(6)

            # Kalem tablosu
            if receipt.items:
                table = self.doc.add_table(rows=len(receipt.items) + 1, cols=2)
                table.style = 'Light Grid Accent 1'

                # Başlık satırı
                header_cells = table.rows[0].cells
                header_cells[0].text = "Ürün / Kalem"
                header_cells[1].text = "Fiyat (₺)"
                
                for cell in header_cells:
                    self._set_cell_background(cell, 'D3D3D3')
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True

                # Kalem satırları
                for idx, item in enumerate(receipt.items, start=1):
                    cells = table.rows[idx].cells
                    cells[0].text = item.name
                    cells[1].text = f"₺ {item.price:,.2f}"
            else:
                self.doc.add_paragraph("Bu fişte kalem bilgisi bulunmamaktadır.").italic = True

            # Boş satır (fişler arası ayrım)
            self.doc.add_paragraph()

    def generate_report(
        self,
        receipts: List[Receipt],
        summary: BatchReportSummary,
        output_path: str
    ) -> bool:
        """
        Tüm raporları oluşturur ve Word dosyasına kaydeder.
        
        Ana rapor oluşturma metodu.
        Sırasıyla:
        1. Başlık sayfası
        2. Genel özet
        3. Kategori dağılımı
        4. Fiş detayları
        
        Args:
            receipts (List[Receipt]): Receipt nesnelerinin listesi
            summary (BatchReportSummary): Rapor özeti
            output_path (str): Çıktı dosyası yolu
            
        Returns:
            bool: Başarı durumu
        """
        try:
            logger.info(f"Rapor oluşturma başladı - Çıktı: {output_path}")

            # Başlık sayfası
            title_heading = self.doc.add_heading(self.title, level=0)
            title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_heading.runs:
                run.font.size = Pt(24)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 128)

            # Alt başlık
            subtitle = self.doc.add_paragraph(
                f"Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}"
            )
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].italic = True

            self.doc.add_paragraph()  # Boş satır

            # İçerik
            self._add_summary_section(summary)
            self.doc.add_page_break()
            
            self._add_category_distribution(summary)
            self._add_receipt_details(receipts)

            # Belgeyi kaydet
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            self.doc.save(str(output_path))
            logger.info(f"Rapor başarıyla kaydedildi: {output_path}")
            return True

        except Exception as e:
            logger.error(f"Rapor oluşturmada hata: {str(e)}")
            return False

    def add_custom_section(self, title: str, content: str) -> None:
        """
        Rapora özel bir bölüm ekler.
        
        Args:
            title (str): Bölüm başlığı
            content (str): Bölüm içeriği
        """
        self._add_heading(title, level=2)
        self.doc.add_paragraph(content)
        logger.debug(f"Özel bölüm eklendi: {title}")

    def add_image(self, image_path: str, width: float = 5.0) -> bool:
        """
        Rapora resim ekler.
        
        Args:
            image_path (str): Resim dosyası yolu
            width (float): Resim genişliği (inç cinsinden)
            
        Returns:
            bool: Başarı durumu
        """
        try:
            if not Path(image_path).exists():
                logger.warning(f"Resim dosyası bulunamadı: {image_path}")
                return False

            self.doc.add_picture(image_path, width=Inches(width))
            logger.debug(f"Resim eklendi: {image_path}")
            return True

        except Exception as e:
            logger.error(f"Resim eklenmesinde hata: {str(e)}")
            return False
